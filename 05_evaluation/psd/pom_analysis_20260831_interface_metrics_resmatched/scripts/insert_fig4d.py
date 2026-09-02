import sys
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = "/tmp/docx_work/v7.docx"
IMG = "/tmp/docx_work/pom_interface_metrics_figure.png"
OUT = "/tmp/docx_work/v7_updated.docx"

doc = Document(SRC)

def force_ltr(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = pPr.find(qn('w:bidi'))
    if bidi is None:
        bidi = OxmlElement('w:bidi')
        pPr.append(bidi)
    bidi.set(qn('w:val'), '0')
    jc = pPr.find(qn('w:jc'))
    if jc is None:
        jc = OxmlElement('w:jc')
        pPr.append(jc)
    jc.set(qn('w:val'), 'left')

def force_run_ltr(run):
    rPr = run._r.get_or_add_rPr()
    rtl = rPr.find(qn('w:rtl'))
    if rtl is None:
        rtl = OxmlElement('w:rtl')
        rPr.append(rtl)
    rtl.set(qn('w:val'), '0')

# Locate the anchor paragraph: last sentence of Table 3's note, right before "Figure 5."
anchor_idx = None
target_idx = None
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t.startswith("Rehovot has no POM class") or ("pore-vs-solid only" in t and "3-phase" in t):
        anchor_idx = i
    if t.startswith("Figure 5."):
        target_idx = i
        break

if anchor_idx is None or target_idx is None:
    print("ANCHOR NOT FOUND", anchor_idx, target_idx)
    sys.exit(1)

print(f"anchor_idx={anchor_idx} target_idx={target_idx}")
print("anchor text:", doc.paragraphs[anchor_idx].text[:120])
print("target text:", doc.paragraphs[target_idx].text[:60])

anchor_p = doc.paragraphs[anchor_idx]._p

def new_paragraph_after(ref_p):
    new_p = OxmlElement('w:p')
    ref_p.addnext(new_p)
    return new_p

# Build in order, each inserted right after the previous, so final order is:
# anchor -> title -> image -> caption -> (existing) Figure 5.

# 1) Title paragraph (bold)
title_p_elm = new_paragraph_after(anchor_p)
from docx.text.paragraph import Paragraph
title_par = Paragraph(title_p_elm, doc.paragraphs[anchor_idx]._parent)
force_ltr(title_par)
run = title_par.add_run("Figure 4d. POM interface metrics (Track B) — resolution-matched")
run.bold = True
force_run_ltr(run)

# 2) Image paragraph
img_p_elm = new_paragraph_after(title_p_elm)
img_par = Paragraph(img_p_elm, doc.paragraphs[anchor_idx]._parent)
force_ltr(img_par)
img_par.alignment = WD_ALIGN_PARAGRAPH.LEFT
img_run = img_par.add_run()
force_run_ltr(img_run)
img_run.add_picture(IMG, width=Inches(6.2))

# 3) Caption paragraph
cap_p_elm = new_paragraph_after(img_p_elm)
cap_par = Paragraph(cap_p_elm, doc.paragraphs[anchor_idx]._parent)
force_ltr(cap_par)
cap_text = (
    "POM interface metrics — Bnei Re’em (Vertisol, n=1) vs. Mishmar HaNegev (Loess, mean ± SE, "
    "n=2, both replicates resolution-matched to ~15 µm). (a) Specific surface area (SSA, POM surface "
    "÷ POM volume) is statistically indistinguishable between soils (30.6±4.4 vs. 33.2 mm²/mm³ total). "
    "(b) Interfacial area density (IAD, interface ÷ bulk sample volume) is 1.6–2.1× higher in Mishmar for "
    "both the POM–pore interface (0.251±0.042 vs. 0.156 mm²/mm³) and the POM–matrix interface "
    "(0.230±0.020 vs. 0.110 mm²/mm³), despite indistinguishable SSA — i.e. the soils differ in how much "
    "POM there is per unit soil volume, not in how convoluted the POM surface itself is. (c) Voxel-face contact-fraction "
    "cross-check (independent method, same direction as a/b). (d) Total POM surface area by two estimation methods "
    "(marching cubes vs. voxel-face counting); the two disagree in absolute mm² (validated against synthetic-sphere "
    "ground truth — marching cubes is closer to the true value) but agree on the pore-vs-matrix split within ~8 "
    "percentage points, which is what (a)–(c) rely on. (e) Interfacial area is more concentrated in a few large "
    "objects in Mishmar (top-5 objects = 54.9±3.8% of total POM–pore interface) than in Bnei Re’em (33.9%). "
    "Methods: POM–pore/POM–matrix contact obtained by 6-connected voxel-face adjacency counting (Schlüter et al., "
    "2014); interfacial surface area independently estimated via marching-cubes mesh reconstruction (Vogel & Roth, 2001; "
    "Houston et al., 2013; Juyal et al., 2021), each triangle assigned to the pore or matrix phase by its outward normal; "
    "IAD normalized by bulk sample volume following the same convention as this document's connectivity-density metric "
    "(Herring et al., 2015; Table 3). n=1 for Bnei Re’em precludes an omnibus significance test; no p-values or "
    "letters are reported, consistent with Table 3's connectivity metrics above."
)
run = cap_par.add_run(cap_text)
run.italic = True
run.font.size = Pt(9)
force_run_ltr(run)

doc.save(OUT)
print("SAVED", OUT)
