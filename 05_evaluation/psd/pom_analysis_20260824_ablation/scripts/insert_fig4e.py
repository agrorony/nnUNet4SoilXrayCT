import sys
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph

SRC = "/tmp/docx_work/v7_before_4e.docx"
IMG = "/tmp/docx_work/pom_size_spatial_figure.png"
OUT = "/tmp/docx_work/v7_with_4e.docx"

doc = Document(SRC)

def force_ltr(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = pPr.find(qn('w:bidi'))
    if bidi is None:
        bidi = OxmlElement('w:bidi'); pPr.append(bidi)
    bidi.set(qn('w:val'), '0')
    jc = pPr.find(qn('w:jc'))
    if jc is None:
        jc = OxmlElement('w:jc'); pPr.append(jc)
    jc.set(qn('w:val'), 'left')

def force_run_ltr(run):
    rPr = run._r.get_or_add_rPr()
    rtl = rPr.find(qn('w:rtl'))
    if rtl is None:
        rtl = OxmlElement('w:rtl'); rPr.append(rtl)
    rtl.set(qn('w:val'), '0')

# Anchor: the last paragraph of the Figure 4d caption (ends with the
# Table 3 letters-convention sentence). Target: the "Figure 5." placeholder.
anchor_idx = None
target_idx = None
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if "consistent with Table 3" in t and "connectivity metrics" in t:
        anchor_idx = i
    if t.startswith("Figure 5."):
        target_idx = i
        break

if anchor_idx is None or target_idx is None:
    print("ANCHOR NOT FOUND", anchor_idx, target_idx)
    sys.exit(1)

print(f"anchor_idx={anchor_idx} target_idx={target_idx}")
print("anchor text:", doc.paragraphs[anchor_idx].text[:120])
print("target text:", doc.paragraphs[target_idx].text[:60])

anchor_p = doc.paragraphs[anchor_idx]._p

def new_paragraph_after(ref_p):
    new_p = OxmlElement('w:p')
    ref_p.addnext(new_p)
    return new_p

# 1) Title paragraph (bold)
title_p_elm = new_paragraph_after(anchor_p)
title_par = Paragraph(title_p_elm, doc.paragraphs[anchor_idx]._parent)
force_ltr(title_par)
run = title_par.add_run("Figure 4e. POM object size distribution and spatial pattern")
run.bold = True
force_run_ltr(run)
title_par.paragraph_format.keep_with_next = True

# 2) Image paragraph
img_p_elm = new_paragraph_after(title_p_elm)
img_par = Paragraph(img_p_elm, doc.paragraphs[anchor_idx]._parent)
force_ltr(img_par)
img_par.alignment = WD_ALIGN_PARAGRAPH.LEFT
img_par.paragraph_format.keep_with_next = True
img_run = img_par.add_run()
force_run_ltr(img_run)
img_run.add_picture(IMG, width=Inches(6.2))

# 3) Caption paragraph
cap_p_elm = new_paragraph_after(img_p_elm)
cap_par = Paragraph(cap_p_elm, doc.paragraphs[anchor_idx]._parent)
force_ltr(cap_par)
cap_text = (
    "POM object size distribution and spatial arrangement — Vertisol (Bnei Re'em, n=1) vs. "
    "Loess (Mishmar HaNegev, n=1), each at its own native analysis resolution (15.0 µm and 5.85 µm "
    "respectively); a legacy result from the original 2-soil POM comparison, unaffected by the "
    "shape/archetype-clustering work that was later retired (2026-08-30) as not robust. "
    "(a) Volume-weighted median POM object diameter is similar between soils (734.0 µm Loess vs. "
    "690.2 µm Vertisol). (b) Despite similar median size, POM volume is far more concentrated in a "
    "single large fragment in Loess (largest object = 45.1% of total denoised POM volume) than in "
    "Vertisol (17.1%) — i.e. Loess's POM population is size-skewed toward one dominant object while "
    "Vertisol's is more evenly distributed across many similarly sized fragments. (c) Both soils show "
    "a spatially aggregated (clustered) arrangement of POM objects relative to complete spatial "
    "randomness (Clark-Evans R=1); Vertisol is more strongly clustered (R=0.640) than Loess (R=0.818). "
    "Methods: object segmentation and denoising as in Table 2/pom_analysis_20260815_light; median "
    "diameter and largest-object share computed on the denoised POM object-size distribution "
    "(part_c_4way_comparison.md); Clark-Evans index computed on POM object centroids "
    "(pom_spatial_pattern_summary_2soil_clean.json), R<1 = aggregated, R=1 = CSR, R>1 = regular/dispersed. "
    "n=1 per soil precludes an omnibus significance test; no p-values or letters are reported, "
    "consistent with Table 3's connectivity metrics above."
)
run = cap_par.add_run(cap_text)
run.italic = True
run.font.size = Pt(9)
force_run_ltr(run)

doc.save(OUT)
print("SAVED", OUT)
